import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# Function to create Word report
def create_word_report(df_summary, chart_buf, pie_chart_buf, total_checks):
    doc = Document()

    doc.add_heading('Báo Cáo Đánh Giá Công Tác Kiểm Tra Áp Giá', 0)

    # Add Summary
    doc.add_heading('1. Tổng quan:', level=1)
    doc.add_paragraph(f'Tổng số lượng kiểm tra toàn công ty: {total_checks:,} lượt.')
    if total_checks == 0:
        doc.add_paragraph("Chưa có dữ liệu kiểm tra.")
    else:
        doc.add_paragraph("Báo cáo thể hiện so sánh giữa số khách hàng có thay đổi và tổng số lượt kiểm tra tại từng điện lực. Tỷ lệ thay đổi được thể hiện trong bảng và biểu đồ bên dưới.")

    # Thêm nhận xét về top 3 và bottom 3 (trước khi đổi tên cột)
    top_3 = df_summary.sort_values(by='Tỷ lệ thay đổi (%)', ascending=False).head(3)
    bottom_3 = df_summary.sort_values(by='Tỷ lệ thay đổi (%)', ascending=True).head(3)

    doc.add_heading('2. Đánh giá nổi bật:', level=1)
    doc.add_paragraph("- Các điện lực có tỷ lệ thay đổi cao nhất:")
    for idx, row in top_3.iterrows():
        doc.add_paragraph(f"  • {row['Area']}: {row['Tỷ lệ thay đổi (%)']}% ({int(row['Total_Changes'])} KH thay đổi trên {int(row['Total_Checks'])} lượt kiểm tra)")

    doc.add_paragraph("- Các điện lực có tỷ lệ thay đổi thấp nhất:")
    for idx, row in bottom_3.iterrows():
        doc.add_paragraph(f"  • {row['Area']}: {row['Tỷ lệ thay đổi (%)']}% ({int(row['Total_Changes'])} KH thay đổi trên {int(row['Total_Checks'])} lượt kiểm tra)")

    # Đổi tên cột
    column_mapping = {
        'Area': 'Điện lực',
        'Total_Checks': 'Tổng số lượng kiểm tra',
        'Total_Changes': 'Số KH có thay đổi'
    }
    df_export = df_summary.rename(columns=column_mapping)

    # Làm tròn số và xử lý NaN
    df_export['Tổng số lượng kiểm tra'] = df_export['Tổng số lượng kiểm tra'].fillna(0).round(0).astype(int)
    df_export['Số KH có thay đổi'] = df_export['Số KH có thay đổi'].fillna(0).round(0).astype(int)

    # Add data table
    doc.add_heading('3. Bảng tổng hợp:', level=1)
    table = doc.add_table(rows=1, cols=len(df_export.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df_export.columns):
        hdr_cells[i].text = column

    for index, row in df_export.iterrows():
        if row['Điện lực'] and str(row['Điện lực']).strip().lower() not in ['nan', 'đơn vị']:
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                if isinstance(item, float) and 'tỷ lệ' in df_export.columns[i].lower():
                    row_cells[i].text = f"{item:.2f}%"
                else:
                    row_cells[i].text = f"{item}"

    # Add Charts
    doc.add_heading('4. Biểu Đồ So Sánh:', level=1)
    doc.add_picture(chart_buf, width=Inches(6))

    doc.add_heading('5. Biểu Đồ Tỷ Lệ thay đổi vs không thay đổi:', level=1)
    doc.add_picture(pie_chart_buf, width=Inches(4.5))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# MAIN APP

st.set_page_config(page_title="Đánh giá kiểm tra áp giá", layout="wide")
st.title("Đánh giá công tác kiểm tra áp giá")

uploaded_file = st.file_uploader("Tải lên file Excel dữ liệu", type=["xlsx"])

if uploaded_file is not None:
    df = pd.read_excel(uploaded_file)
    df.columns = df.columns.str.strip()
    df.columns = ['Stt', 'Area', 'Check_SH_2plus', 'Check_HCSN', 'Check_Production', 'Check_KDDV', 'Check_PriceRate', 'Check_SH_Level3', 'Total_Checks', 'Total_Changes']

    df_summary = df[['Area', 'Total_Checks', 'Total_Changes']].copy()
    df_summary['Total_Checks'] = pd.to_numeric(df_summary['Total_Checks'], errors='coerce')
    df_summary['Total_Changes'] = pd.to_numeric(df_summary['Total_Changes'], errors='coerce')
    df_summary = df_summary[df_summary['Area'].notna()]
    df_summary['Area'] = df_summary['Area'].astype(str)
    df_summary = df_summary[df_summary['Area'].str.strip().str.lower() != 'tổng cộng']

    total_checks = df_summary['Total_Checks'].sum()
    total_changes = df_summary['Total_Changes'].sum()
    df_summary['Tỷ lệ thay đổi (%)'] = (df_summary['Total_Changes'] / df_summary['Total_Checks'] * 100).round(2)

    # Biểu đồ cột đôi
    fig, ax = plt.subplots(figsize=(12,6))
    x = df_summary['Area']
    bar_width = 0.35
    index = range(len(x))

    ax.bar(index, df_summary['Total_Checks'], bar_width, label='Tổng số lượng kiểm tra')
    ax.bar([i + bar_width for i in index], df_summary['Total_Changes'], bar_width, label='Số KH có thay đổi')

    ax.set_xlabel('Điện lực')
    ax.set_ylabel('Số lượng')
    ax.set_title('So sánh giữa Tổng số kiểm tra và Số KH có thay đổi')
    ax.set_xticks([i + bar_width / 2 for i in index])
    ax.set_xticklabels(x, rotation=90)
    ax.legend()
    plt.tight_layout()

    chart_buf = BytesIO()
    fig.savefig(chart_buf)
    chart_buf.seek(0)
    st.pyplot(fig)

    # Pie chart
    fig_pie, ax_pie = plt.subplots()
    ax_pie.pie([total_changes, total_checks - total_changes], labels=['Có thay đổi', 'Không thay đổi'], autopct='%1.1f%%', colors=['#ff9999','#66b3ff'])
    ax_pie.set_title('Tỷ lệ KH có thay đổi vs không thay đổi')
    pie_chart_buf = BytesIO()
    fig_pie.savefig(pie_chart_buf)
    pie_chart_buf.seek(0)
    st.pyplot(fig_pie)

    st.dataframe(df_summary)

    today = datetime.today().strftime('%Y-%m-%d')
    word_file = create_word_report(df_summary, chart_buf, pie_chart_buf, total_checks)

    st.download_button(
        label="📄 Tải báo cáo Word",
        data=word_file,
        file_name=f'Bao_cao_kiem_tra_ap_gia_{today}.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
